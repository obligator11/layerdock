"""
DOCX generation engine.

Consumes the structured dict produced by pdf_parser.parse_pdf() and builds
a real .docx: text lines as positioned textboxes, images/graphics as
positioned pictures, and now TABLES as actual editable Word table objects
— floated to their exact page position via OOXML's tblpPr, since
python-docx has no built-in support for absolutely positioned tables.
"""
import io
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Emu, Pt
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from backend.pdf_parser import parse_pdf

EMU_PER_PT = 12700
_next_id = [1]


def _pt_to_emu(pt):
    return int(pt * EMU_PER_PT)


def _pt_to_twips(pt):
    return int(pt * 20)


def _rgb_hex(rgb):
    return "%02X%02X%02X" % tuple(rgb)


def _embed_image_get_relid(paragraph, img_bytes):
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(img_bytes))
    drawing = run._element.find(qn('w:drawing'))
    blip = drawing.find('.//' + qn('a:blip'))
    rel_id = blip.get(qn('r:embed'))
    paragraph._p.remove(run._element)
    return rel_id


def _run_xml(run):
    half_pts = max(int(run["size"] * 2), 2)
    color = _rgb_hex(run["color"])
    safe_text = escape(run["text"])
    safe_font = escape(run["font"] or "Calibri")
    font_lower = (run["font"] or "").lower()
    bold = "bold" in font_lower or bool(run.get("flags", 0) & 16)
    italic = "italic" in font_lower or "oblique" in font_lower or bool(run.get("flags", 0) & 2)

    b_tag = "<w:b/>" if bold else ""
    i_tag = "<w:i/>" if italic else ""

    return f"""
      <w:r>
        <w:rPr>
          <w:rFonts w:ascii="{safe_font}" w:hAnsi="{safe_font}"/>
          {b_tag}{i_tag}
          <w:sz w:val="{half_pts}"/>
          <w:color w:val="{color}"/>
        </w:rPr>
        <w:t xml:space="preserve">{safe_text}</w:t>
      </w:r>
    """


def _line_textbox_xml(x_pt, y_pt, w_pt, h_pt, runs):
    _next_id[0] += 1
    shape_id = _next_id[0]
    x, y = _pt_to_emu(x_pt), _pt_to_emu(y_pt)
    cx, cy = max(_pt_to_emu(w_pt), 1), max(_pt_to_emu(h_pt), 1)

    runs_xml = "".join(_run_xml(r) for r in runs)

    xml = f"""
    <w:r {nsdecls('w')}>
      <w:drawing xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="{shape_id}" behindDoc="0" locked="0"
                   layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="Line{shape_id}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:cNvSpPr txBox="1"/>
                <wps:spPr>
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                  <a:ln><a:noFill/></a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p>
                      <w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
                      {runs_xml}
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr wrap="none" lIns="0" tIns="0" rIns="0" bIns="0" anchor="t">
                  <a:noAutofit/>
                </wps:bodyPr>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </w:r>
    """
    return parse_xml(xml)


def _picture_anchor_xml(x_pt, y_pt, w_pt, h_pt, rel_id):
    _next_id[0] += 1
    shape_id = _next_id[0]
    x, y = _pt_to_emu(x_pt), _pt_to_emu(y_pt)
    cx, cy = max(_pt_to_emu(w_pt), 1), max(_pt_to_emu(h_pt), 1)
    rel_id = escape(str(rel_id))

    xml = f"""
    <w:r {nsdecls('w')}>
      <w:drawing xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="{shape_id}" behindDoc="0" locked="0"
                   layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>
          <wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="Picture{shape_id}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:nvPicPr>
                  <pic:cNvPr id="{shape_id}" name="Picture{shape_id}"/>
                  <pic:cNvPicPr/>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{rel_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                  <a:stretch><a:fillRect/></a:stretch>
                </pic:blipFill>
                <pic:spPr>
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </w:r>
    """
    return parse_xml(xml)


def _add_floating_table(doc, table_data):
    """Build a real, editable Word table and float it to an exact page
    position via w:tblpPr — OOXML's absolute table-positioning element,
    which python-docx doesn't expose directly, hence the raw XML."""
    rows, cols = table_data["rows"], table_data["cols"]
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    tbl = table._tbl
    tblPr = tbl.tblPr

    tblpPr = OxmlElement("w:tblpPr")
    tblpPr.set(qn("w:vertAnchor"), "page")
    tblpPr.set(qn("w:horzAnchor"), "page")
    tblpPr.set(qn("w:tblpX"), str(_pt_to_twips(table_data["bbox"][0])))
    tblpPr.set(qn("w:tblpY"), str(_pt_to_twips(table_data["bbox"][1])))
    tblPr.append(tblpPr)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    total_width_pt = sum(table_data["col_widths"])
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_pt_to_twips(total_width_pt)))
    tblPr.append(tblW)

    tblGrid = tbl.find(qn("w:tblGrid"))
    grid_cols = tblGrid.findall(qn("w:gridCol"))
    for gc, w in zip(grid_cols, table_data["col_widths"]):
        gc.set(qn("w:w"), str(_pt_to_twips(w)))

    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), str(_pt_to_twips(max(table_data["row_heights"][r_idx], 10))))
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)

        for c_idx, cell in enumerate(row.cells):
            cell.text = table_data["cells"][r_idx][c_idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(_pt_to_twips(table_data["col_widths"][c_idx])))
            tcPr.append(tcW)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def build_docx(pdf_path: str, output_path: str, progress_cb=None) -> dict:
    _next_id[0] = 1
    parsed = parse_pdf(pdf_path)
    doc = Document()
    flagged_pages = []

    for i, page in enumerate(parsed["pages"]):
        conf = page.get("ocr_confidence")
        if conf is not None and conf < 60:
            flagged_pages.append(page["number"])

        section = doc.sections[0] if i == 0 else doc.add_section()
        section.page_width = Emu(_pt_to_emu(page["width"]))
        section.page_height = Emu(_pt_to_emu(page["height"]))
        section.left_margin = section.right_margin = Emu(0)
        section.top_margin = section.bottom_margin = Emu(0)

        anchor_paragraph = doc.add_paragraph()

        for img in page["images"]:
            try:
                rel_id = _embed_image_get_relid(anchor_paragraph, img["raw_bytes"])
                x0, y0, x1, y1 = img["bbox"]
                run_elem = _picture_anchor_xml(x0, y0, x1 - x0, y1 - y0, rel_id)
                anchor_paragraph._p.append(run_elem)
            except Exception as e:
                import traceback
                print(f"[docx_builder] FAILED to embed image on page {page['number']}: {e}")
                traceback.print_exc()
                continue

        for line in page["text_lines"]:
            x0, y0, x1, y1 = line["bbox"]
            max_size = max((r["size"] for r in line["runs"]), default=10)
            h = max(y1 - y0, max_size * 1.25)
            run_elem = _line_textbox_xml(x0, y0, max(x1 - x0, 4), h, line["runs"])
            anchor_paragraph._p.append(run_elem)

        for table_data in page.get("tables", []):
            try:
                _add_floating_table(doc, table_data)
            except Exception as e:
                import traceback
                print(f"[docx_builder] FAILED to build table on page {page['number']}: {e}")
                traceback.print_exc()
                continue

        if progress_cb:
            progress_cb(i + 1, parsed["page_count"])

    doc.save(output_path)
    return {
        "page_count": parsed["page_count"],
        "output_path": output_path,
        "flagged_pages": flagged_pages,
    }